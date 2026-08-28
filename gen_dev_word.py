"""
V2.5.0 图文版 Word PRD - 完全重构（方案 A）
- §二 拆 12 子节（§1 预诊 ... §12 危急值）
- 每个子节末尾插入对应截图
- 表格真渲染（python-docx Table）
- §三 改附录索引（图编号 + 跳回引用）
"""
import re
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========== 资源路径 ==========
PRD = Path('/Users/zdhome/Documents/工作内容/仲超/Obsidian Vault/02-AI产品经理/作品集/PRD/AI辅助问诊_PRD优化版_V2.0.md')
SHOTS = Path('/tmp/demo_shots_v2/v230_pages')
OUT = Path('/Users/zdhome/Documents/工作内容/仲超/Obsidian Vault/02-AI产品经理/作品集/PRD/AI辅助问诊_PRD_V2.5.0_图文版_开发交付.docx')

# ========== 路由→截图 映射（每张图带章节引用）==========
ROUTE_SHOTS = [
    # (截图文件名, 标题, [插入到的章节列表])
    ('01_login',         '图 1-1：登录页（手机号+验证码）',                          ['§1', '§2']),
    ('02_pre_diagnosis', '图 1-2：预诊入口（手机壳+大屏一体机）',                    ['§1']),
    ('03_pre_diagnosis_chat', '图 1-3：预诊对话满屏（AI 引导 5 轮）',                ['§1']),
    ('04_register',      '图 2-1：挂号页（患者到达 → 生成诊疗号）',                  ['§2']),
    ('05_dashboard',     '图 5-1：医生工作台（4 入口 + 高风险红条）',                ['§5']),
    ('06_emr_new',       '图 6-1：P3 接诊中（左 7 流程 + 右 5 AI·展开）',          ['§4', '§6', '§7']),
    ('07_prescription',  '图 8-1：处方页（开方弹窗 + 过敏拦截）',                    ['§8']),
    ('08_critical_1',    '图 12-1：危急值列表',                                      ['§12']),
    ('09_followup_config', '图 9-1：随访任务配置（5 任务）',                          ['§9']),
    ('10_followup',      '图 9-2：随访工作台（风险降序）',                            ['§9', '§10']),
    ('11_followup_patient_1', '图 10-1：风险患者详情（4 类风险）',                    ['§10']),
    ('12_stats',         '图 11-1：运营统计（5 指标）',                                ['§11']),
    ('13_emr_leftpanel_closed', '图 6-2：P3 接诊中（收起 PRD 抽屉）',                  ['§6']),
    ('14_dashboard_leftpanel_closed', '图 5-2：工作台（收起态）',                     ['§5']),
    ('15_pre_chat_full', '图 1-4：预诊对话满屏（移动端）',                            ['§1']),
    ('16_critical_4ch',  '图 12-2：危急值 4 通道详情',                                ['§12']),
]

# 章节→对应截图（便于 §X 末尾插图）
CHAPTER_SHOTS = {
    '§1': ['01_login', '02_pre_diagnosis', '03_pre_diagnosis_chat', '15_pre_chat_full'],
    '§2': ['01_login', '04_register'],
    '§4': ['06_emr_new'],
    '§5': ['05_dashboard', '14_dashboard_leftpanel_closed'],
    '§6': ['06_emr_new', '13_emr_leftpanel_closed'],
    '§7': ['06_emr_new'],
    '§8': ['07_prescription'],
    '§9': ['09_followup_config', '10_followup'],
    '§10': ['11_followup_patient_1', '10_followup'],
    '§11': ['12_stats'],
    '§12': ['08_critical_1', '16_critical_4ch'],
}

# 章节标题（§1-§12）
CHAPTER_TITLES = {
    '§1': '模块 1：预诊',
    '§2': '模块 2：档案',
    '§3': '模块 3：录音',
    '§4': '模块 4：诊中分流',
    '§5': '模块 5：检查回流',
    '§6': '模块 6：AI 诊断',
    '§7': '模块 7：病历',
    '§8': '模块 8：处方',
    '§9': '模块 9：智能随访',
    '§10': '模块 10：AI 风险识别',
    '§11': '模块 11：多租户',
    '§12': '模块 12：危急值',
}

# ========== 字体工具 ==========
def set_cn_font(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Arial')

# ========== 文档元素工具 ==========
def add_title(doc, text, size=24, color=RGBColor(0x1A, 0x1A, 0x2E), align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=True, color=color)
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_cn_font(run, size=18, bold=True, color=RGBColor(0x7C, 0x3A, 0xED))
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_cn_font(run, size=14, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_cn_font(run, size=12, bold=True, color=RGBColor(0x37, 0x41, 0x51))
    return p

def add_body(doc, text, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_cn_font(run, size=10, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_cn_font(run, size=9)
            run.font.name = 'Courier New'
        elif part:
            run = p.add_run(part)
            set_cn_font(run, size=10)
    return p

def add_li(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_cn_font(run, size=10, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_cn_font(run, size=9)
            run.font.name = 'Courier New'
        elif part:
            run = p.add_run(part)
            set_cn_font(run, size=10)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_cn_font(run, size=9, color=RGBColor(0x6B, 0x72, 0x80))
    return p

def add_image(doc, path, width_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p

def add_image_caption(doc, text, bold_label=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bold_label:
        run = p.add_run(bold_label)
        set_cn_font(run, size=9, bold=True, color=RGBColor(0x6B, 0x72, 0x80))
    run = p.add_run(text)
    set_cn_font(run, size=9, color=RGBColor(0x6B, 0x72, 0x80))
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_ncol(doc, header, rows, col_widths_cm=None, header_color='7C3AED'):
    """真表格渲染（python-docx Table）"""
    ncols = len(header)
    t = doc.add_table(rows=1, cols=ncols)
    t.style = 'Light Grid Accent 1'
    # 表头
    for i, h in enumerate(header):
        cell = t.cell(0, i)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cell.paragraphs[0].add_run(h)
        set_cn_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_color)
        tcPr.append(shd)
    # 数据行：逐行 add_row（安全避免越界）
    for row in rows:
        new_row = t.add_row()
        # 补齐列数
        while len(new_row.cells) < ncols:
            new_row.add_cell()
        # 截断多余
        for c_idx in range(ncols):
            if c_idx < len(row):
                val = str(row[c_idx])
            else:
                val = ''
            cell = new_row.cells[c_idx]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(val)
            set_cn_font(run, size=9)
    if col_widths_cm:
        for r in t.rows:
            for c_idx, w in enumerate(col_widths_cm):
                if c_idx < len(r.cells):
                    r.cells[c_idx].width = Cm(w)
    return t

def add_page_break(doc):
    doc.add_page_break()

# ========== PRD 章节拆分（核心逻辑）==========
def split_prd_to_chapters(prd_text):
    """解析 PRD md → 返回 12 章节 dict（key=§X, value=markdown content）"""
    lines = prd_text.split('\n')
    chapters = {}
    current_key = '__preface__'  # 章节前的导言
    chapters[current_key] = []

    for line in lines:
        # 匹配 §X（一级 # 二级 ## 三级 ### 标题里的"§X"）
        # PRD V2.0 用 "## 模块 1：预诊" / "## 一、xxx" 等格式
        m = re.match(r'^#{1,3}\s+(?:模块\s*)?(\d+)[：:、\.]?\s*(.*)', line)
        if m:
            num = m.group(1)
            # 1-12 章节
            if num in [str(i) for i in range(1, 13)]:
                new_key = f'§{num}'
                if new_key not in chapters:
                    chapters[new_key] = []
                current_key = new_key
                chapters[current_key].append(line)
                continue
        # §X.Y 章节
        m2 = re.match(r'^#{2,4}\s+([一二三四五六七八九十]+)、(.+)', line)
        if m2:
            # 中文数字章节（如"一、xxx"）也归到对应数字 §X
            cn_to_num = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10, '十一': 11, '十二': 12}
            num_cn = m2.group(1)
            if num_cn in cn_to_num:
                num = str(cn_to_num[num_cn])
                if num in [str(i) for i in range(1, 13)]:
                    new_key = f'§{num}'
                    if new_key not in chapters:
                        chapters[new_key] = []
                    current_key = new_key
        chapters[current_key].append(line)
    return chapters

# ========== 渲染单个章节（含表格+插图）==========
def render_chapter(doc, chapter_key, chapter_lines, screenshots):
    """渲染单章节：标题 + 内容 + 末尾插图"""
    if chapter_key.startswith('§'):
        num = chapter_key[1:]
        title = f'§{num} {CHAPTER_TITLES.get(chapter_key, "")}'
        add_h1(doc, title)
    else:
        # 导言
        pass

    # 渲染内容
    in_table = False
    table_buf = []
    in_code = False
    code_buf = []
    for line in chapter_lines:
        # 代码块
        if line.startswith('```'):
            if in_code:
                add_quote(doc, '\n'.join(code_buf[:20]))
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        # 表格检测
        if line.strip().startswith('|') and not in_table:
            in_table = True
            table_buf = [line]
            continue
        if in_table:
            if line.strip().startswith('|'):
                table_buf.append(line)
                continue
            else:
                # 渲染表格
                render_table(doc, table_buf)
                in_table = False
                table_buf = []

        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level, txt = len(m.group(1)), m.group(2).strip()
            if level <= 2:
                add_h2(doc, txt)
            elif level == 3:
                add_h3(doc, txt)
            else:
                add_h3(doc, txt)
            continue
        # 列表
        if re.match(r'^[-*]\s+', line):
            txt = re.sub(r'^[-*]\s+', '', line).strip()
            add_li(doc, txt)
            continue
        # 引用
        if line.startswith('>'):
            txt = line.lstrip('>').strip()
            add_quote(doc, txt)
            continue
        # 分隔
        if line.strip() in ('---', '***'):
            continue
        if not line.strip():
            continue
        # 段落
        add_body(doc, line.strip())

    # 末尾插图
    if screenshots:
        add_h2(doc, '📷 对应 demo 截图')
        for fname in screenshots:
            shot_path = SHOTS / f'{fname}.png'
            if not shot_path.exists():
                continue
            # 找标题
            title = next((t for f, t, _ in ROUTE_SHOTS if f == fname), fname)
            add_image(doc, shot_path, width_cm=14)
            add_image_caption(doc, title.replace('：', ' · '))

def render_table(doc, table_buf):
    """渲染 markdown 表格为真 python-docx Table"""
    if len(table_buf) < 2:
        return
    # 解析表头
    header_line = table_buf[0]
    headers = [c.strip() for c in header_line.split('|')[1:-1]]
    if not headers or all(not h for h in headers):
        return
    # 解析数据行（跳过分隔行 |---|）
    rows = []
    for row_line in table_buf[2:]:
        cells = [c.strip() for c in row_line.split('|')[1:-1]]
        # 简化 markdown
        clean_cells = []
        for c in cells:
            # 去 **bold** 标记（保留文字）
            c = re.sub(r'\*\*([^*]+)\*\*', r'\1', c)
            c = re.sub(r'`([^`]+)`', r'\1', c)
            # 转义 & 符号
            c = c.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            clean_cells.append(c)
        if clean_cells:
            rows.append(clean_cells)
    if not rows:
        return
    # 等宽列
    ncols = len(headers)
    col_w = 15.0 / ncols
    add_table_ncol(doc, headers, rows, col_widths_cm=[col_w]*ncols)

# ========== 主构建 ==========
def build():
    print(f'开始生成 Word V2.5.0: {OUT.name}')
    prd_text = PRD.read_text(encoding='utf-8')
    chapters = split_prd_to_chapters(prd_text)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ===== 封面 =====
    doc.add_paragraph()
    add_title(doc, 'AI 辅助问诊系统', size=28)
    add_title(doc, 'PRD V2.5.0 · 图文版 · 开发交付文档', size=16,
              color=RGBColor(0x7C, 0x3A, 0xED))
    doc.add_paragraph()
    add_body(doc, '产品经理：仲超')
    add_body(doc, '最后更新：2026-08-28（V2.5.0）')
    add_body(doc, '配套 demo：https://csebk.github.io/portfolio/ai-diagnosis-demo/')
    doc.add_paragraph()
    add_body(doc, '核心保证：')
    add_li(doc, '✅ 开发读完不需要再问任何问题')
    add_page_break(doc)

    # ===== 目录 =====
    add_title(doc, '目  录', size=22)
    doc.add_paragraph()
    toc_rows = [
        ('§一', '文档说明 & 双向追溯'),
        ('§二', 'PRD 12 大模块（每节含 demo 截图）'),
        ('§二.1', '模块 1：预诊（4 截图）'),
        ('§二.2', '模块 2：档案（2 截图）'),
        ('§二.4', '模块 4：诊中分流（1 截图）'),
        ('§二.5', '模块 5：检查回流（2 截图）'),
        ('§二.6', '模块 6：AI 诊断（2 截图）'),
        ('§二.7', '模块 7：病历（1 截图）'),
        ('§二.8', '模块 8：处方（1 截图）'),
        ('§二.9', '模块 9：智能随访（2 截图）'),
        ('§二.10', '模块 10：AI 风险识别（2 截图）'),
        ('§二.11', '模块 11：多租户（1 截图）'),
        ('§二.12', '模块 12：危急值（2 截图）'),
        ('§三', '附录：截图索引（16 张）'),
        ('§四', '验收标准 29 条 + 上线 checklist'),
        ('§五', '附录：开发上手 + 链接 + 术语'),
    ]
    add_table_ncol(doc, ['章节', '内容'], toc_rows, col_widths_cm=[2, 13])
    add_page_break(doc)

    # ===== §一 文档说明 =====
    add_h1(doc, '§一、文档说明 & 双向追溯')
    add_h2(doc, '1.1 本文档面向开发')
    add_body(doc, '本文档是 AI 辅助问诊系统的**完整产品需求**，配套 demo（链接见封面）。'
                '§二 12 大模块**每节末尾**插入对应 demo 截图——'
                '开发读 §2.6 时，AI 诊断截图就在旁边，**无需前后翻对照**。')
    add_h2(doc, '1.2 demo ↔ PRD 双向追溯')
    add_body(doc, '本系统的 PRD 文本与 demo 数据通过 prd_data.json（单一真相源）双向绑定：')
    add_li(doc, '改 prd_data.json 任意字段 → demo 立即更新（XHR 同步加载）')
    add_li(doc, '改 demo 源码 → 跑 sync_prd.js demo-default 反向抽数据到 JSON')
    add_li(doc, '**改 prd_data.json → 跑 sync_prd.js rules-by-page** → 重新生成 §X.7 6 字段汇总')
    add_li(doc, '**开发修改 demo 时，优先改 JSON**，避免硬编码漂移')
    add_h2(doc, '1.3 关键设计决策')
    add_table_ncol(doc, ['决策', '当前规格', '原因'],
        [
            ['页面数', '16', '诊前+诊中+诊后三段式'],
            ['分栏', '左 7 + 右 5', '医生一气呵成不跳页'],
            ['录音', '1 次开/1 次关', '医生双手操作电脑反馈'],
            ['号源', '释放（不锁）', '基层号源紧张'],
            ['检查结果', 'AI 二次解读', '差异化卖点'],
            ['危机值', '4 通道强推', '临床安全底线'],
            ['30min 已读', '强制+升级', '《医疗事故处理条例》'],
            ['状态显示', '12 状态码', '诊疗流程状态（非在线/离线）'],
            ['PRD 规则树', '按页过滤 60 条', '开发自查不出错'],
            ['Tip 提示', '该字段 6 条规则表格', 'hover 即查即用'],
        ],
        col_widths_cm=[3, 4, 6])
    add_page_break(doc)

    # ===== §二 PRD 12 大模块（每节末尾插图）=====
    add_h1(doc, '§二、PRD 12 大模块（每节含 demo 截图）')
    add_body(doc, '本节按 §1 预诊 → §12 危急值 顺序展示 12 大模块。'
                '**每节末尾**插入对应 demo 截图（来自 GitHub Pages V2.5.0 远程）。')
    add_quote(doc, '说明：§3 录音 / §2.3 / §2.4 中部分内容嵌入对应大章节，未单独截图。')

    # 按 §1-§12 顺序渲染
    for ch_num in range(1, 13):
        ch_key = f'§{ch_num}'
        if ch_key not in chapters:
            continue
        add_page_break(doc)
        add_h1(doc, f'§二.{ch_num} {CHAPTER_TITLES.get(ch_key, "")}')
        shots = CHAPTER_SHOTS.get(ch_key, [])
        render_chapter(doc, ch_key, chapters[ch_key], shots)

    # ===== §三 附录：截图索引 =====
    add_page_break(doc)
    add_h1(doc, '§三、附录：截图索引（16 张）')
    add_body(doc, '本节为**附录索引**，所有 demo 截图已在 §二 12 大模块中按章节穿插展示。'
                '此处提供完整索引（截图编号 + 路由 + 关键交互）便于快速定位。')
    add_quote(doc, '提示：开发查阅时，按章节回到 §二.X 即可看到对应截图与详细规则。')

    add_h2(doc, '3.1 完整截图索引表')
    add_table_ncol(doc, ['图编号', '截图文件', '路由', '关键交互', '对应 §X'],
        [
            ['图 1-1', '01_login', '#/login', '手机号+验证码（中国用户）', '§1 §2'],
            ['图 1-2', '02_pre_diagnosis', '#/pre-diagnosis', '预诊入口（手机壳+大屏）', '§1'],
            ['图 1-3', '03_pre_diagnosis_chat', '#/pre-diagnosis/chat', 'AI 引导 5 轮对话', '§1'],
            ['图 1-4', '15_pre_chat_full', '#/pre-diagnosis/chat', '预诊满屏（移动端）', '§1'],
            ['图 2-1', '04_register', '#/register', '挂号→刷身份证→诊疗号', '§2'],
            ['图 5-1', '05_dashboard', '#/dashboard', '4 入口+高风险+最近患者', '§5'],
            ['图 5-2', '14_dashboard_leftpanel_closed', '#/dashboard', '工作台收起 PRD 抽屉', '§5'],
            ['图 6-1', '06_emr_new', '#/emr/new', '接诊中 左 7+右 5 展开', '§4 §6 §7'],
            ['图 6-2', '13_emr_leftpanel_closed', '#/emr/new', '接诊中 收起 PRD 抽屉', '§6'],
            ['图 8-1', '07_prescription', '#/prescription', '开方弹窗+过敏拦截', '§8'],
            ['图 9-1', '09_followup_config', '#/followup-config', '5 任务配置', '§9'],
            ['图 9-2', '10_followup', '#/followup', '随访工作台（风险降序）', '§9 §10'],
            ['图 10-1', '11_followup_patient_1', '#/followup/patient/1', '风险患者详情', '§10'],
            ['图 11-1', '12_stats', '#/stats', '5 指标统计', '§11'],
            ['图 12-1', '08_critical_1', '#/critical/1', '危急值列表', '§12'],
            ['图 12-2', '16_critical_4ch', '#/critical/1', '4 通道详情', '§12'],
        ],
        col_widths_cm=[1.5, 3, 2.5, 4, 2])

    # ===== §四 验收标准 =====
    add_page_break(doc)
    add_h1(doc, '§四、验收标准 29 条 + 上线 checklist')

    add_h2(doc, '4.1 功能验收 AC1-6')
    for ac in [
        ('AC1', '预诊建议准确率 ≥ 70%（基于医生采纳率）'),
        ('AC2', '录音转写 30s 内完成（mock 1.8s）'),
        ('AC3', '接诊中页面 0 跳转（V2.0.1 一页式）'),
        ('AC4', '过敏冲突 100% 拦截（青霉素+阿莫西林→红色警告+不保存）'),
        ('AC5', 'AI 诊断采纳/忽略/编辑→日志可查（含原 AI 输出）'),
        ('AC6', '诊所 A 数据绝不进诊所 B（多租户 tenant_id 强制）'),
    ]:
        add_li(doc, f'**{ac[0]}**：{ac[1]}')

    add_h2(doc, '4.2 V2.5.0 新增功能 AC7-11')
    for ac in [
        ('AC7', '诊中分流 4 边界场景：①24h 超时转异常 ②危急值 4 通道强推+30min 已读 ③换医生→按开单医生推送+值班医生接管 ④医生下班→队列保留 7 天+同事接管'),
        ('AC8', '健康档案 opt-in：患者主动授权→病历自动带入（不授权为空）'),
        ('AC9', '检查结果 AI 二次解读 V2.0 必做（病历+检查联合推理）'),
        ('AC10', '录音完成→右栏"主诉/现病史"待确认→医生点"接受"AI 语义合并到左栏（草稿保留 ✅ 标记）'),
        ('AC11', 'PRD 抽屉展开时 14 Tip + "💡 PRD 追溯"可见；收起时全部隐藏'),
        ('AC12', 'V2.5.0 PRD 规则树按 currentPath 过滤，**当前页只显示该页规则**'),
        ('AC13', 'V2.5.0 Tip hover 显示**该字段全部 6 条详细规则**（表格）'),
    ]:
        add_li(doc, f'**{ac[0]}**：{ac[1]}')

    add_h2(doc, '4.3 危急值验收 AC14-17')
    for ac in [
        ('AC14', '4 通道强推：工作台红点（立即）+ App push（30s 内）+ 短信（60s 内）+ 大屏（同步）'),
        ('AC15', '30min 内已读强制（未读→升级科室主任；60min 仍未读→系统电话呼叫）'),
        ('AC16', '3 级升级链路：医生→主任（30min）→院长（60min）'),
        ('AC17', '100% 审计留痕（每次推送/已读/升级/处理全记录）'),
    ]:
        add_li(doc, f'**{ac[0]}**：{ac[1]}')

    add_h2(doc, '4.4 业务验收 AB1-6')
    for ac in [
        ('AB1', 'AI 诊断采纳率 ≥ 60%'),
        ('AB2', '单患者接诊时间 < 8 分钟（不含检查等待）'),
        ('AB3', 'AI 误诊率 < 8%'),
        ('AB4', '复诊率 +20%（随访推送后）'),
        ('AB5', '患者响应率 ≥ 70%（随访任务）'),
        ('AB6', '医生满意度 ≥ 4.0/5.0（季度调研）'),
    ]:
        add_li(doc, f'**{ac[0]}**：{ac[1]}')

    add_h2(doc, '4.5 性能/安全/合规 AP/AS/AR')
    add_li(doc, '**AP1-AP5**：预诊 API <2s / AI 推理 <3s / 病历保存 <1s / 危急值推送 <1s / 100 医生并发')
    add_li(doc, '**AS1-AS4**：HTTPS + AES-256 / 越权=0（tenant_id 强制）/ 审计日志 5 年 / 防爆破 5 次/15min')
    add_li(doc, '**AR1-AR4**：《医疗事故处理条例》4 通道 / 《互联网医疗管理办法》实名+同意 / 《个保法》opt-in / 《电子病历规范》CA 签名')

    add_h2(doc, '4.6 上线 checklist（10 条）')
    for ck in [
        '✅ PRD V2.5.0 全部 12 模块评审通过',
        '✅ 29 条验收标准全部自测通过',
        '✅ 5 大异常流程演练（24h 超时/危急值/换医生/下班/AI 宕机）',
        '✅ 5 大边界场景演练（断网/重连/重复点击/状态冲突/历史数据迁移）',
        '✅ 多租户隔离测试（诊所 A 数据不可见诊所 B）',
        '✅ 危急值 30min 已读压力测试（100 条同时推送）',
        '✅ 性能压测（100 医生 × 8 小时 = 800 接诊）',
        '✅ 安全渗透测试（SQL 注入/XSS/CSRF/越权）',
        '✅ 合规预审（法务/医疗顾问/网信办备案）',
        '✅ 培训 5 角色（超管/院长/医生/护士/患者代表）',
    ]:
        add_li(doc, ck)

    # ===== §五 附录 =====
    add_page_break(doc)
    add_h1(doc, '§五、附录')

    add_h2(doc, '5.1 开发上手 3 步')
    for step in [
        '**第 1 步：**打开 demo → https://csebk.github.io/portfolio/ai-diagnosis-demo/ → 浏览 16 路由',
        '**第 2 步：**读 PRD V2.5.0（§二 12 模块）→ 重点看每个模块的 §X.7 6 字段汇总 + 末尾截图',
        '**第 3 步：**按 §四 29 条验收标准逐条自测 → 完成后通知 PM 验收',
    ]:
        add_li(doc, step)

    add_h2(doc, '5.2 关键链接')
    for k, v in [
        ('Demo 永久地址', 'https://csebk.github.io/portfolio/ai-diagnosis-demo/'),
        ('Demo 源码仓库', 'https://github.com/csebk/portfolio/tree/main/ai-diagnosis-demo'),
        ('PRD V2.5.0 md', 'https://github.com/csebk/portfolio/blob/main/AI辅助问诊_PRD优化版_V2.0.md'),
        ('JSON 单一真相源', 'https://github.com/csebk/portfolio/blob/main/ai-diagnosis-demo/js/prd_data.json'),
        ('sync_prd.js 双向同步', 'https://github.com/csebk/portfolio/blob/main/sync_prd.js'),
        ('rules-by-page 命令', 'node sync_prd.js rules-by-page'),
    ]:
        add_li(doc, f'**{k}**：{v}')

    add_h2(doc, '5.3 术语表')
    add_table_ncol(doc, ['术语', '释义'],
        [
            ['PRD', 'Product Requirements Document，产品需求文档'],
            ['V2.5.0', '当前版本（V2.0 demo 自 0.1 → 2.5.0 共 14 个迭代）'],
            ['诊中分流', 'V2.0 核心创新：开检查单→释放号源→结果回流→AI 二次解读'],
            ['4 通道强推', '工作台红点 + App push + 短信 + 大屏'],
            ['12 状态码', 'pre_diagnosis/registered/waiting/consulting/test_ordered/critical_value/prescribed/followup_configured/followup_active/completed/missed/awaiting_test'],
            ['tenant_id', '多租户隔离字段（每张表必带）'],
            ['AI 二次解读', 'V2.0 差异化卖点：检查结果回流后 AI 联合病历+检查推理'],
            ['rules_by_page', 'V2.5.0 新增：按页分组的详细规则（60 条 / 7 页面）'],
            ['6 字段汇总', 'PRD 每个模块的 §X.7 表格：触发/结束/状态/权限/异常/验收'],
        ],
        col_widths_cm=[3, 12])

    # 保存
    doc.save(str(OUT))
    size_kb = OUT.stat().st_size // 1024
    print(f'✅ 生成成功：{OUT} ({size_kb}KB)')


if __name__ == '__main__':
    build()
